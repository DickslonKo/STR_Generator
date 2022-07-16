import pandas as pd
from docx import Document
from docx.shared import Inches
from datetime import datetime
import numpy as np
from datetime import date
from selenium import webdriver
from webdriver_manager.chrome import ChromeDriverManager
from selenium.webdriver.chrome.service import Service
from selenium.webdriver import ChromeOptions
from bs4 import BeautifulSoup
import streamlit as st
from google.cloud import storage
import os
import time

st.title('Welcome to Suspicious Transaction Report Generator')

st.text("Designed and Developed by Dickson Ko")

uploaded_files = st.file_uploader("Please choose an alert_assessment_form", type=["csv"])

uploaded_files1 = st.file_uploader("Please choose a transaction history", type=["csv"])

if uploaded_files is not None:
    df = pd.read_csv(uploaded_files, header=None)

if uploaded_files1 is not None:
    df_txn = pd.read_csv(uploaded_files1)

now = datetime.now()
now_name = now.strftime("%H%M%S")

bucket_name = 'final_project0820'
STR_name = f'STR_{now_name}.docx'
internet_search_name = f'Internet_search_{now_name}.docx'
os.environ.setdefault("My First Project", "quantum-engine-346005")
storage_client = storage.Client.from_service_account_json("quantum-engine-346005-76af3461d45b.json")
bucket = storage_client.get_bucket(bucket_name)


if st.button('Generate a STR'):
    if df.loc[12][2] is not np.nan:
        ind_name = df.loc[0][2]
        DOB_obj = datetime.strptime(df.loc[11][2], '%d-%b-%y')
        age = date.today().year - DOB_obj.year - ((date.today().month, date.today().day) < (DOB_obj.month, DOB_obj.day))
        nationality = ' '.join(map(str, df.loc[17][3].split(" ")[2:]))
        occupation = df.loc[12][2]
        employer = df.loc[15][2]
        establishment = df.loc[19][2]
    else:
        pass

    if df.loc[23][2] is not np.nan:
        Com_name = df.loc[0][2]
        nature = df.loc[23][2]
        incor_country = df.loc[26][3]
        address = df.loc[25][2]
        open_date = df.loc[28][2]
        roles = df.loc[30][2].split("\n")
        key_person = df.loc[30][3].split("\n")
    else:
        pass

    pair = {}

    if df.loc[23][2] is not np.nan:
        for i in range(len(key_person)):
            if key_person[i] in pair:
                if not isinstance(pair[key_person[i]], list):
                    pair[key_person[i]] = [pair[key_person[i]]]
                pair[key_person[i]].append(roles[i])
            else:
                pair[key_person[i]]=[roles[i]]
    else:
        pass

    UBO = ""
    for i in pair:
        UBO += f"{i} is "
        if len(pair[i]) == 1:
            UBO += f"{pair[i][0]}. "
        if len(pair[i]) == 2:
            for j in range(len(pair[i])):
                if j == 1:
                    UBO += f"{pair[i][j]}. "
                else:
                    UBO += f"{pair[i][j]} and "
        if len(pair[i]) == 3:
            for j in range(len(pair[i])):
                if j == 2:
                    UBO += f"{pair[i][j]}. "
                elif j == 0:
                    UBO += f"{pair[i][j]},"
                else:
                    UBO += f"{pair[i][j]} and "
        if len(pair[i]) == 4:
            for j in range(len(pair[i])):
                if j == 3:
                    UBO += f"{pair[i][j]}. "
                elif j == 0 or 1:
                    UBO += f"{pair[i][j]},"
                else:
                    UBO += f"{pair[i][j]} and "

    df_txn["Date"] = pd.to_datetime(df_txn["Date"])
    df_txn["Date"] = df_txn["Date"].dt.date
    df_txn = df_txn.set_index("Date")
    df_txn["Amount in Cr (HKD)"] = pd.to_numeric(df_txn["Amount in Cr (HKD)"])
    df_txn["Amount in Dr (HKD)"] = pd.to_numeric(df_txn["Amount in Dr (HKD)"])

    cr_sum = df_txn["Amount in Cr (HKD)"].sum()
    cr_count = df_txn["Amount in Cr (HKD)"].count()
    dr_sum = df_txn["Amount in Dr (HKD)"].sum()
    dr_count = df_txn["Amount in Dr (HKD)"].count()

    Cr_tx_type = df_txn.groupby("Transaction type")["Amount in Cr (HKD)"].sum()
    Cr_tx_type = Cr_tx_type.drop(Cr_tx_type[Cr_tx_type.values ==0].index)

    Cr_type_count = df_txn.groupby("Transaction type")["Amount in Cr (HKD)"].count()
    Cr_type_count = Cr_type_count.drop(Cr_type_count[Cr_type_count.values ==0].index)

    Cr_type_order = []
    for i in np.argsort(Cr_tx_type)[::-1][:3]:
        Cr_type_order.append(i)


    Cr_Cpty = df_txn.groupby("Counterparty name")["Amount in Cr (HKD)"].sum()
    Cr_Cpty = Cr_Cpty.drop(Cr_Cpty[Cr_Cpty.values ==0].index)

    Cr_Cpty_no = df_txn.groupby("Counterparty name")["Amount in Cr (HKD)"].count()
    Cr_Cpty_no = Cr_Cpty_no.drop(Cr_Cpty_no[Cr_Cpty_no.values ==0].index)

    Cr_Cpty_order = []
    for i in np.argsort(Cr_Cpty)[::-1][:5]:
        Cr_Cpty_order.append(i)

    Dr_tx_type = df_txn.groupby("Transaction type")["Amount in Dr (HKD)"].sum()
    Dr_tx_type = Dr_tx_type.drop(Dr_tx_type[Dr_tx_type.values ==0].index)

    Dr_type_count = df_txn.groupby("Transaction type")["Amount in Dr (HKD)"].count()
    Dr_type_count = Dr_type_count.drop(Dr_type_count[Dr_type_count.values ==0].index)

    Dr_type_order = []
    for i in np.argsort(Dr_tx_type)[::-1][:3]:
        Dr_type_order.append(i)

    Dr_Cpty = df_txn.groupby("Counterparty name")["Amount in Dr (HKD)"].sum()
    Dr_Cpty = Dr_Cpty.drop(Dr_Cpty[Dr_Cpty.values ==0].index)

    Dr_Cpty_no = df_txn.groupby("Counterparty name")["Amount in Dr (HKD)"].count()
    Dr_Cpty_no = Dr_Cpty_no.drop(Dr_Cpty_no[Dr_Cpty_no.values ==0].index)

    Dr_Cpty_order = []
    for i in np.argsort(Dr_Cpty)[::-1][:5]:
        Dr_Cpty_order.append(i)

    cr_type = ""

    for i in Cr_type_order:
        if i == Cr_type_order[-1]:
            cr_type += f" and HKD {Cr_tx_type.values[i]} was {Cr_tx_type.index[i]}({Cr_type_count[i]} counts)."
        else:
            cr_type += f" HKD {Cr_tx_type.values[i]} was {Cr_tx_type.index[i]}({Cr_type_count[i]} counts),"

    dr_type = ""

    for i in Dr_type_order:
        if i == Dr_type_order[-1]:
            dr_type += f" and HKD {Dr_tx_type.values[i]} was {Dr_tx_type.index[i]}({Dr_type_count[i]} counts)."
        else:
            dr_type += f" HKD {Dr_tx_type.values[i]} was {Dr_tx_type.index[i]}({Dr_type_count[i]} counts),"

    document = Document()

    if df.loc[12][2] is not np.nan:
        introduction2 = document.add_paragraph(f"{ind_name}, aged {age}, is {nationality} citizen and declared to be {occupation} of {employer}. {ind_name} has banked with us since {establishment}.")
    if df.loc[23][2] is not np.nan:
        introduction = document.add_paragraph(f"{Com_name}, incorporated in {incor_country}, is declared to be {nature} and has banked with us since {open_date}.")
        introduction1 = document.add_paragraph(f"The business address of {Com_name} is {address}.")
        introduction2 = document.add_paragraph(f"{UBO}")
    run = introduction2.add_run()
    run.add_break()

    txn_review0 = document.add_paragraph(f"Account activities from {df_txn.index.min()} to {df_txn.index.max()} have been reviewed.")
    txn_review1 = document.add_paragraph(f"There were {cr_count} mixed credit transactions totaling HKD {cr_sum} of which {cr_type} The incoming fund was mainly from following counterparties:")
    for i in Cr_Cpty_order:
        if i == Cr_Cpty_order[-1]:
            document.add_paragraph(f"HKD {Cr_Cpty[Cr_Cpty_order[i]]} was made from {Cr_Cpty.index[Cr_Cpty_order[i]]}({Cr_Cpty_no[Cr_Cpty_order[i]]} counts).", style="List Bullet")
        else:
            document.add_paragraph(f"HKD {Cr_Cpty[Cr_Cpty_order[i]]} was made from {Cr_Cpty.index[Cr_Cpty_order[i]]}({Cr_Cpty_no[Cr_Cpty_order[i]]} counts);", style="List Bullet")
    txn_review2 = document.add_paragraph(f"There were {dr_count} mixed debit transactions totaling HKD {dr_sum} of which {dr_type} The outgoing fund was mainly made to following counterparties:")
    for i in Dr_Cpty_order:
        if i == Dr_Cpty_order[-1]:
            document.add_paragraph(f"HKD {Dr_Cpty[i]} was made to {Dr_Cpty.index[i]}({Dr_Cpty_no[i]} counts).", style="List Bullet")
        else:
            document.add_paragraph(f"HKD {Dr_Cpty[i]} was made to {Dr_Cpty.index[i]}({Dr_Cpty_no[i]} counts);", style="List Bullet")

    document.save(r"STR.docx")

    blob = bucket.blob(STR_name)
    blob.upload_from_filename(f'./STR.docx')
    options = webdriver.ChromeOptions()
    options.add_argument('--headless')
    options.add_argument('--disable-gpu')
    options.add_argument('--no-sandbox')
    options.add_argument('--disable-dev-shm-usage')
    driver = webdriver.Chrome(ChromeDriverManager().install(),options=options)
    driver.implicitly_wait(0.5)
    query = list(Cr_Cpty.index) + list(Dr_Cpty.index)
    document1 = Document()

    for i in range(len(query)):
        p = document1.add_paragraph()
        r = p.add_run()
        r.add_text(query[i])
        url = "http://www.google.com/search?q=" + query[i] + "&start=" +      'str((page - 1) * 3)'
        driver.get(url)
        google_search = driver.save_screenshot("google.png")
        r.add_picture("google.png", width=Inches(6), height=Inches(3.8))
        soup = BeautifulSoup(driver.page_source, 'html.parser')

        search = soup.find_all('div', class_="yuRUbf")
        for h in search[:3]:
            try:
                driver.get(h.a.get('href'))
                image = driver.save_screenshot('image.png')
                r.add_picture('image.png', width=Inches(6), height=Inches(3.8))
                r.add_text(h.a.get('href'))
            except:
                pass
            continue

    document1.save(r"Internet search.docx")

    blob = bucket.blob(internet_search_name)
    blob.upload_from_filename(f'./Internet search.docx')

    driver.quit()
    st.write("Job Completed! Your STR report and internet search result can be downloaded via following links.")
    st.subheader("STR Report")
    st.markdown(f'https://storage.googleapis.com/{bucket_name}/{STR_name}')
    st.subheader("Internet Search")
    st.markdown(f'https://storage.googleapis.com/{bucket_name}/{internet_search_name}')

